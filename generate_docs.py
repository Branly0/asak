from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Title
title = doc.add_heading('API Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Test Management System')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.italic = True

doc.add_paragraph()

# Base URL Section
doc.add_heading('Base URL', level=1)
doc.add_paragraph('http://localhost:8000')

doc.add_heading('Authentication', level=1)
auth_intro = doc.add_paragraph('All endpoints require JWT Bearer token authentication except /auth/register and /auth/login')
doc.add_paragraph('Include token in header:', style='List Bullet')
code_block = doc.add_paragraph('Authorization: Bearer <access_token>')
code_block_format = code_block.paragraph_format
code_block_format.left_indent = Inches(0.5)

# ============ AUTH ENDPOINTS ============
doc.add_page_break()
doc.add_heading('Authentication Endpoints', level=1)

# Register
doc.add_heading('1. Register User', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Property'
hdr_cells[1].text = 'Value'

rows_data = [
    ('Method', 'POST'),
    ('Endpoint', '/auth/register'),
    ('Auth', 'None'),
    ('Content-Type', 'application/json'),
]
for key, val in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = val

doc.add_heading('Request Body', level=3)
code = doc.add_paragraph('''{
  "name": "John Doe",
  "age": 25,
  "sex": "male",
  "email": "john@example.com",
  "password": "secure_password",
  "role": "evaluator"  // or "pupil"
}''')
code.style = 'No Spacing'
code_format = code.paragraph_format
code_format.left_indent = Inches(0.5)

doc.add_heading('Response (200)', level=3)
code = doc.add_paragraph('''{
  "id": "550e8400-e29b-41d4-a716-446655440000",
  "email": "john@example.com",
  "name": "John Doe",
  "age": 25
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

# Login
doc.add_heading('2. Login', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Property'
hdr_cells[1].text = 'Value'

rows_data = [
    ('Method', 'POST'),
    ('Endpoint', '/auth/login'),
    ('Auth', 'None'),
    ('Content-Type', 'application/x-www-form-urlencoded'),
]
for key, val in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = val

doc.add_heading('Request Body', level=3)
code = doc.add_paragraph('''username=john@example.com&password=secure_password''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('Response (200)', level=3)
code = doc.add_paragraph('''{
  "access_token": "eyJhbGciOiJIUzI1NiIs...",
  "refresh_token": "eyJhbGciOiJIUzI1NiIs...",
  "token_type": "bearer"
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

# Refresh Token
doc.add_heading('3. Refresh Token', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Property'
hdr_cells[1].text = 'Value'

rows_data = [
    ('Method', 'POST'),
    ('Endpoint', '/auth/refresh'),
    ('Auth', 'None'),
]
for key, val in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = val

doc.add_heading('Request Body', level=3)
code = doc.add_paragraph('''{"refresh_token": "eyJhbGciOiJIUzI1NiIs..."}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('Response (200)', level=3)
code = doc.add_paragraph('''{
  "access_token": "eyJhbGciOiJIUzI1NiIs...",
  "token_type": "bearer"
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

# Logout
doc.add_heading('4. Logout', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Property'
hdr_cells[1].text = 'Value'

rows_data = [
    ('Method', 'POST'),
    ('Endpoint', '/auth/logout'),
    ('Auth', 'Required'),
]
for key, val in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = val

doc.add_heading('Response (200)', level=3)
code = doc.add_paragraph('''{"message": "Successfully logged out"}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

# ============ ASSESSMENT ENDPOINTS ============
doc.add_page_break()
doc.add_heading('Assessment Endpoints', level=1)

# Teacher endpoints section
doc.add_heading('TEACHER ENDPOINTS', level=2)

# Create Test
doc.add_heading('1. Create Test', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Property'
hdr_cells[1].text = 'Value'

rows_data = [
    ('Method', 'POST'),
    ('Endpoint', '/tests/create'),
    ('Auth', 'Required (Teacher only)'),
    ('Content-Type', 'application/json'),
]
for key, val in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = val

doc.add_heading('Request Body', level=3)
code = doc.add_paragraph('''{
  "name": "Math Quiz 2024",
  "description": "Basic algebra questions"
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('Response (200)', level=3)
code = doc.add_paragraph('''{
  "id": "550e8400-e29b-41d4-a716-446655440000",
  "name": "Math Quiz 2024",
  "description": "Basic algebra questions",
  "is_published": false,
  "evaluator_id": "550e8400-e29b-41d4-a716-446655440001",
  "created_at": "2024-06-12T10:30:00Z",
  "updated_at": "2024-06-12T10:30:00Z"
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

# Add Questions
doc.add_heading('2. Add Questions to Test', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Property'
hdr_cells[1].text = 'Value'

rows_data = [
    ('Method', 'POST'),
    ('Endpoint', '/tests/{test_id}/questions'),
    ('Auth', 'Required (Test owner)'),
    ('Content-Type', 'application/json'),
]
for key, val in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = val

doc.add_heading('Request Body', level=3)
code = doc.add_paragraph('''{
  "question_text": "What is 2 + 2?",
  "question_type": "multiple_choice",
  "question_number": 1,
  "answers": [
    {"answer_text": "4", "is_correct": true},
    {"answer_text": "3", "is_correct": false},
    {"answer_text": "5", "is_correct": false}
  ]
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('Question Types', level=3)
doc.add_paragraph('multiple_choice', style='List Bullet')
doc.add_paragraph('true_false', style='List Bullet')
doc.add_paragraph('short_answer', style='List Bullet')

# Upload PDF
doc.add_heading('3. Upload PDF & Extract Questions', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Property'
hdr_cells[1].text = 'Value'

rows_data = [
    ('Method', 'POST'),
    ('Endpoint', '/tests/{test_id}/upload-pdf'),
    ('Auth', 'Required (Test owner)'),
    ('Content-Type', 'multipart/form-data'),
]
for key, val in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = val

doc.add_heading('Query Parameters', level=3)
doc.add_paragraph('past_exam (boolean, optional): true for past exam papers, false for generating Q&A', style='List Bullet')

doc.add_heading('Request Body', level=3)
code = doc.add_paragraph('''file: <PDF file>
past_exam: false''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('Response (200)', level=3)
code = doc.add_paragraph('''{
  "message": "Extracted 5 questions from PDF",
  "test_id": "550e8400-e29b-41d4-a716-446655440000",
  "questions_count": 5
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

# Publish Test
doc.add_heading('4. Publish Test', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Property'
hdr_cells[1].text = 'Value'

rows_data = [
    ('Method', 'PATCH'),
    ('Endpoint', '/tests/{test_id}/publish'),
    ('Auth', 'Required (Test owner)'),
]
for key, val in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = val

doc.add_heading('Response (200)', level=3)
code = doc.add_paragraph('''{
  "id": "550e8400-e29b-41d4-a716-446655440000",
  "name": "Math Quiz 2024",
  "is_published": true,
  ...
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

# View Results
doc.add_heading('5. View Test Results', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Property'
hdr_cells[1].text = 'Value'

rows_data = [
    ('Method', 'GET'),
    ('Endpoint', '/tests/{test_id}/results'),
    ('Auth', 'Required (Test owner)'),
]
for key, val in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = val

doc.add_heading('Response (200)', level=3)
code = doc.add_paragraph('''{
  "test_id": "550e8400-e29b-41d4-a716-446655440000",
  "test_name": "Math Quiz 2024",
  "results": [
    {
      "student_id": "550e8400-e29b-41d4-a716-446655440002",
      "student_name": "Alice",
      "student_email": "alice@example.com",
      "score": 80,
      "total_questions": 5,
      "submitted_at": "2024-06-12T11:30:00Z"
    }
  ]
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

# Student endpoints section
doc.add_heading('STUDENT ENDPOINTS', level=2)

# Get Available Tests
doc.add_heading('6. Get Available Tests', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Property'
hdr_cells[1].text = 'Value'

rows_data = [
    ('Method', 'GET'),
    ('Endpoint', '/tests/available'),
    ('Auth', 'Required (Student only)'),
]
for key, val in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = val

doc.add_heading('Response (200)', level=3)
code = doc.add_paragraph('''{
  [
    {
      "id": "550e8400-e29b-41d4-a716-446655440000",
      "name": "Math Quiz 2024",
      "description": "Basic algebra",
      "is_published": true,
      ...
    }
  ]
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

# Get Test Questions
doc.add_heading('7. Get Test Questions', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Property'
hdr_cells[1].text = 'Value'

rows_data = [
    ('Method', 'GET'),
    ('Endpoint', '/tests/{test_id}/questions'),
    ('Auth', 'Required (Student only)'),
]
for key, val in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = val

doc.add_heading('Response (200)', level=3)
code = doc.add_paragraph('''{
  "id": "550e8400-e29b-41d4-a716-446655440000",
  "name": "Math Quiz 2024",
  "questions": [
    {
      "id": "550e8400-e29b-41d4-a716-446655440003",
      "question_text": "What is 2 + 2?",
      "question_type": "multiple_choice",
      "question_number": 1,
      "answers": [
        {"id": "550e8400...", "answer_number": 1, "answer_text": "4", "is_correct": false},
        {"id": "550e8400...", "answer_number": 2, "answer_text": "3", "is_correct": false},
        {"id": "550e8400...", "answer_number": 3, "answer_text": "5", "is_correct": false}
      ]
    }
  ]
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

# Submit Test
doc.add_heading('8. Submit Test Answers', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Property'
hdr_cells[1].text = 'Value'

rows_data = [
    ('Method', 'POST'),
    ('Endpoint', '/tests/{test_id}/submit'),
    ('Auth', 'Required (Student only)'),
    ('Content-Type', 'application/json'),
]
for key, val in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = val

doc.add_heading('Request Body', level=3)
code = doc.add_paragraph('''{
  "answers": [
    {
      "question_id": "550e8400-e29b-41d4-a716-446655440003",
      "selected_answer_id": "550e8400-e29b-41d4-a716-446655440005"
    },
    {
      "question_id": "550e8400-e29b-41d4-a716-446655440004",
      "answer_text": "My short answer"
    }
  ]
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('Response (200) - Includes Correct Answers', level=3)
code = doc.add_paragraph('''{
  "id": "550e8400-e29b-41d4-a716-446655440006",
  "test_id": "550e8400-e29b-41d4-a716-446655440000",
  "pupil_id": "550e8400-e29b-41d4-a716-446655440002",
  "score": 80,
  "total_questions": 5,
  "submitted_at": "2024-06-12T12:00:00Z",
  "student_answers": [
    {
      "id": "550e8400...",
      "question_id": "550e8400...",
      "selected_answer_id": "550e8400...",
      "is_correct": true,
      "correct_answer_text": "4",
      ...
    }
  ]
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

# View Result Detail
doc.add_heading('9. View Result Detail', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Property'
hdr_cells[1].text = 'Value'

rows_data = [
    ('Method', 'GET'),
    ('Endpoint', '/results/{result_id}'),
    ('Auth', 'Required'),
]
for key, val in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = val

doc.add_heading('Response (200)', level=3)
code = doc.add_paragraph('''{
  "id": "550e8400-e29b-41d4-a716-446655440006",
  "test_id": "550e8400-e29b-41d4-a716-446655440000",
  "pupil_id": "550e8400-e29b-41d4-a716-446655440002",
  "score": 80,
  "total_questions": 5,
  "submitted_at": "2024-06-12T12:00:00Z",
  "student_answers": [
    {
      "id": "550e8400...",
      "question_id": "550e8400...",
      "question_text": "What is 2 + 2?",
      "selected_answer_id": "550e8400...",
      "answer_text": null,
      "is_correct": true,
      "correct_answer_id": "550e8400...",
      "correct_answer_text": "4"
    }
  ]
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

# Error Responses
doc.add_page_break()
doc.add_heading('Error Responses', level=1)

doc.add_heading('401 - Unauthorized', level=2)
code = doc.add_paragraph('''{
  "detail": "Invalid authentication credentials"
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('403 - Forbidden', level=2)
code = doc.add_paragraph('''{
  "detail": "Only teachers can perform this action"
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('404 - Not Found', level=2)
code = doc.add_paragraph('''{
  "detail": "Test not found"
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

doc.add_heading('400 - Bad Request', level=2)
code = doc.add_paragraph('''{
  "detail": "Cannot add questions to published tests"
}''')
code.style = 'No Spacing'
code.paragraph_format.left_indent = Inches(0.5)

# Save
doc.save('/home/nkasi/Documents/work/asak/API_Documentation.docx')
print("✅ API documentation saved: API_Documentation.docx")
